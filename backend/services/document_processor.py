from __future__ import annotations

from dataclasses import asdict, dataclass
from io import BytesIO
from pathlib import Path
from typing import Callable, Dict

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from fpdf import FPDF

from attacks.transformers import load_pdf_document, PdfDocument


class DocumentProcessingError(ValueError):
    """Raised when an uploaded document cannot be converted into a PDF."""


@dataclass
class DocumentMetadata:
    source_name: str
    source_extension: str
    content_type: str
    char_length: int
    page_count: int | None = None
    converter: str | None = None

    def to_dict(self) -> Dict[str, object]:
        return asdict(self)


@dataclass
class ProcessedDocument:
    pdf_bytes: bytes
    pdf_document: PdfDocument
    metadata: DocumentMetadata


_TEXT_EXTENSIONS = {".txt", ".md", ".rst"}


class DocumentProcessor:
    """Convert arbitrary supported documents into PDF payloads."""

    def __init__(self) -> None:
        self._handlers: Dict[str, Callable[[bytes, str], ProcessedDocument]] = {
            ".pdf": self._process_pdf,
            ".docx": self._process_docx,
            ".html": self._process_html,
            ".htm": self._process_html,
        }
        for ext in _TEXT_EXTENSIONS:
            self._handlers[ext] = self._process_text

    def process(self, data: bytes, filename: str) -> ProcessedDocument:
        if not data:
            raise DocumentProcessingError("The uploaded file is empty.")

        suffix = Path(filename).suffix.lower()
        handler = self._handlers.get(suffix)
        if handler is None:
            raise DocumentProcessingError(
                f"Unsupported file extension '{suffix or 'unknown'}'. Allowed: PDF, DOCX, HTML, TXT, MD."
            )
        return handler(data, filename)

    # -- individual handlers -------------------------------------------------

    def _process_pdf(self, data: bytes, filename: str) -> ProcessedDocument:
        document = load_pdf_document(data, filename)
        metadata = DocumentMetadata(
            source_name=filename,
            source_extension=Path(filename).suffix.lower() or ".pdf",
            content_type="application/pdf",
            char_length=len(document.text),
            page_count=document.page_count,
            converter=None,
        )
        return ProcessedDocument(pdf_bytes=data, pdf_document=document, metadata=metadata)

    def _process_docx(self, data: bytes, filename: str) -> ProcessedDocument:
        try:
            doc = DocxDocument(BytesIO(data))
        except Exception as exc:
            raise DocumentProcessingError("Could not read the DOCX document.") from exc

        parts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        text = "\n\n".join(parts)
        if not text:
            raise DocumentProcessingError("The DOCX file does not contain extractable text.")
        return self._text_payload_to_pdf(text, filename, converter="docx->pdf", content_type="application/docx")

    def _process_html(self, data: bytes, filename: str) -> ProcessedDocument:
        soup = BeautifulSoup(data, "html.parser")
        text = soup.get_text("\n").strip()
        if not text:
            raise DocumentProcessingError("The HTML document does not contain extractable text.")
        return self._text_payload_to_pdf(text, filename, converter="html->pdf", content_type="text/html")

    def _process_text(self, data: bytes, filename: str) -> ProcessedDocument:
        text = data.decode("utf-8", errors="ignore").strip()
        if not text:
            raise DocumentProcessingError("The text document is empty after decoding.")
        content_type = "text/plain"
        return self._text_payload_to_pdf(text, filename, converter="text->pdf", content_type=content_type)

    # -- helpers -------------------------------------------------------------

    def _text_payload_to_pdf(
        self,
        text: str,
        filename: str,
        *,
        converter: str,
        content_type: str,
    ) -> ProcessedDocument:
        pdf_bytes = self._build_pdf_from_text(text)
        new_name = f"{Path(filename).stem}.pdf"
        document = load_pdf_document(pdf_bytes, new_name)
        metadata = DocumentMetadata(
            source_name=filename,
            source_extension=Path(filename).suffix.lower(),
            content_type=content_type,
            char_length=len(text),
            page_count=document.page_count,
            converter=converter,
        )
        return ProcessedDocument(pdf_bytes=pdf_bytes, pdf_document=document, metadata=metadata)

    def _build_pdf_from_text(self, text: str) -> bytes:
        pdf = FPDF()
        pdf.set_auto_page_break(True, margin=15)
        pdf.add_page()
        self._set_pdf_font(pdf)
        line_height = 6
        usable_width = pdf.w - 2 * pdf.l_margin
        for block in text.splitlines():
            sanitized = block if block.strip() else " "
            pdf.multi_cell(usable_width, line_height, sanitized)
        output = pdf.output(dest="S")
        if isinstance(output, str):
            return output.encode("latin-1", errors="ignore")
        return output

    @staticmethod
    def _set_pdf_font(pdf: FPDF) -> None:
        candidates = [
            ("DejaVu", Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")),
            ("FreeSans", Path("/usr/share/fonts/truetype/freefont/FreeSans.ttf")),
        ]
        for family, font_path in candidates:
            if font_path.exists():
                pdf.add_font(family, "", str(font_path), uni=True)
                pdf.set_font(family, size=11)
                return
        pdf.set_font("Helvetica", size=11)


default_document_processor = DocumentProcessor()

def process_document(data: bytes, filename: str) -> ProcessedDocument:
    return default_document_processor.process(data, filename)
